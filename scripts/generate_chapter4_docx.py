"""
Generate Chapter 4 Word document from Markdown.
Requires: pip install python-docx
Run from project root: python scripts/generate_chapter4_docx.py
Output: docs/Chapter_4_System_Analysis_and_Design.docx
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Install python-docx: pip install python-docx")
    sys.exit(1)

PROJECT_ROOT = Path(__file__).resolve().parent.parent
MD_PATH = PROJECT_ROOT / "docs" / "Chapter 4 System Analysis and Design.md"
OUT_PATH = PROJECT_ROOT / "docs" / "Chapter_4_System_Analysis_and_Design.docx"


def add_paragraph(doc, text, style=None):
    if not text.strip():
        return
    p = doc.add_paragraph(text.strip(), style=style)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text.strip(), level=level)


def add_code_block(doc, lines, lang=""):
    if not lines:
        return
    text = "\n".join(lines)
    p = doc.add_paragraph()
    p.style = "Normal"
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line or not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if not cells:
            continue
        if all(set(c.replace("-", "").replace(" ", "")) <= set(":") for c in cells):
            continue  # skip separator row
        rows.append(cells)
    return rows


def main():
    if not MD_PATH.exists():
        print(f"Not found: {MD_PATH}")
        sys.exit(1)

    doc = Document()

    content = MD_PATH.read_text(encoding="utf-8")
    lines = content.split("\n")
    i = 0
    in_code = False
    code_lang = ""
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        raw = line

        if in_code:
            if line.strip().startswith("```"):
                in_code = False
                add_code_block(doc, code_lines, code_lang)
                code_lines = []
                i += 1
                continue
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            in_code = True
            code_lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            continue

        if in_table:
            if line.strip().startswith("|"):
                table_lines.append(line)
                i += 1
                continue
            else:
                in_table = False
                rows = parse_table(table_lines)
                if rows:
                    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    t.style = "Table Grid"
                    for ri, row in enumerate(rows):
                        for ci, cell in enumerate(row):
                            if ci < len(t.rows[ri].cells):
                                t.rows[ri].cells[ci].text = cell
                table_lines = []
                continue

        if line.strip().startswith("|") and "|" in line[1:]:
            in_table = True
            table_lines = [line]
            i += 1
            continue

        if line.startswith("# "):
            add_heading(doc, line[2:], 0)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 1)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 2)
        elif line.startswith("#### "):
            add_heading(doc, line[5:], 3)
        elif line.strip() == "---":
            doc.add_paragraph()
        else:
            add_paragraph(doc, line)
        i += 1

    if table_lines:
        rows = parse_table(table_lines)
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    if ci < len(t.rows[ri].cells):
                        t.rows[ri].cells[ci].text = cell

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
